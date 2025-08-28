#!/usr/bin/env python3
"""
Script para exportar o glossário para PDF, Word e TXT atualizados
"""

import os
import yaml
import re
from pathlib import Path
from datetime import datetime
import subprocess
import sys

def install_packages():
    """Install required packages if not available"""
    try:
        import markdown
        from docx import Document
    except ImportError:
        print("📦 Instalando dependências...")
        subprocess.check_call([sys.executable, "-m", "pip", "install", 
                              "markdown", "python-docx"])
        import markdown
        from docx import Document

def load_docs_content():
    """Load all documentation content"""
    docs_dir = Path('docs')
    content = {}
    
    # Categories mapping
    categories = {
        'conceitos-fundamentais': 'Conceitos Fundamentais',
        'ia-generativa': 'IA Generativa',
        'agentes-ia': 'Agentes de IA',
        'escopo-ias': 'Escopo das IAs',
        'etica-seguranca-governanca': 'Ética, Segurança e Governança',
        'habilidades-praticas': 'Habilidades e Práticas',
        'infraestrutura-processos': 'Infraestrutura e Processos'
    }
    
    for category_dir, category_name in categories.items():
        category_path = docs_dir / category_dir
        if category_path.exists():
            content[category_name] = []
            
            # Get all .md files except index
            for md_file in sorted(category_path.glob('*.md')):
                if md_file.name != 'index.md':
                    with open(md_file, 'r', encoding='utf-8') as f:
                        file_content = f.read()
                    
                    # Extract title and content
                    lines = file_content.split('\n')
                    title_pt = lines[0].replace('# ', '') if lines else md_file.stem
                    
                    # Extract English term and acronyms from the next non-empty line
                    english_term = ""
                    acronyms = ""
                    
                    for line in lines[1:]:
                        line = line.strip()
                        if not line:
                            continue
                        # Look for **English Term** · *Acronyms* pattern
                        if line.startswith('**') and '**' in line:
                            # Extract English term
                            english_match = re.search(r'\*\*(.*?)\*\*', line)
                            if english_match:
                                english_term = english_match.group(1)
                            
                            # Extract acronyms if present
                            if '·' in line and '*' in line:
                                acronym_match = re.search(r'·\s*\*(.*?)\*', line)
                                if acronym_match:
                                    acronyms = acronym_match.group(1)
                            break
                    
                    # Combine title with English term and acronyms
                    if english_term:
                        if acronyms:
                            title = f"{title_pt} ({english_term} - {acronyms})"
                        else:
                            title = f"{title_pt} ({english_term})"
                    else:
                        title = title_pt
                    
                    # Clean content - remove tags, navigation, emojis, English term line
                    cleaned_lines = []
                    skip_section = False
                    found_english_line = False
                    
                    for line in lines[1:]:
                        # Skip English term line
                        if not found_english_line and line.strip().startswith('**') and '**' in line.strip():
                            found_english_line = True
                            continue
                        # Skip empty lines after English term
                        if found_english_line and not line.strip():
                            found_english_line = False
                            continue
                        # Skip tags line
                        if line.startswith('**Tags:**'):
                            continue
                        # Skip navigation section
                        if line.startswith('---'):
                            skip_section = True
                            continue
                        if skip_section:
                            continue
                        # Remove emoji icons from markdown links
                        line = re.sub(r':\w+:', '', line)  # Remove :material-*: icons
                        # Keep the line if it's content
                        if line.strip():
                            cleaned_lines.append(line)
                    
                    # Get definition (usually first meaningful paragraph)
                    body = '\n'.join(cleaned_lines).strip()
                    
                    # Split into definition and explanation
                    paragraphs = [p.strip() for p in body.split('\n\n') if p.strip()]
                    definition = paragraphs[0] if paragraphs else ''
                    explanation = '\n\n'.join(paragraphs[1:]) if len(paragraphs) > 1 else ''
                    
                    content[category_name].append({
                        'title': title,
                        'definition': definition,
                        'explanation': explanation,
                        'id': md_file.stem
                    })
    
    return content

def build_term_map(content):
    """Build a map of file IDs to term titles for internal linking"""
    term_map = {}
    
    for category_name, terms in content.items():
        for term in terms:
            # Use the file ID (stem) as the key
            file_id = term['id']
            term_map[file_id] = {
                'title': term['title'],
                'category': category_name
            }
    
    return term_map

def process_internal_links(text, term_map):
    """Convert markdown links to internal HTML anchors"""
    # Pattern to match [text](../category/file.md) links
    link_pattern = r'\[([^\]]+)\]\(\.\./[^/]+/([^.]+)\.md\)'
    
    def replace_link(match):
        link_text = match.group(1)
        file_id = match.group(2)
        
        # Check if this file ID exists in our term map
        if file_id in term_map:
            return f'<a href="#{file_id}">{link_text}</a>'
        else:
            # Keep original text if no match found
            return link_text
    
    return re.sub(link_pattern, replace_link, text)

def generate_html(content):
    """Generate HTML from content"""
    install_packages()
    import markdown
    
    # Build term map for internal linking
    term_map = build_term_map(content)
    
    html_parts = ["""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8">
        <title>Glossário de Inteligência Artificial</title>
        <style>
            body { font-family: Arial, sans-serif; line-height: 1.6; margin: 40px; }
            h1 { color: #2c3e50; border-bottom: 3px solid #3498db; padding-bottom: 10px; }
            h2 { color: #34495e; border-bottom: 1px solid #bdc3c7; padding-bottom: 5px; margin-top: 40px; }
            h3 { color: #2980b9; margin-top: 30px; }
            .term { margin-bottom: 30px; border-left: 4px solid #3498db; padding-left: 20px; }
            .meta { color: #7f8c8d; font-size: 0.9em; }
            .toc { background: #f8f9fa; padding: 20px; margin: 20px 0; border-radius: 5px; }
            .toc ul { list-style-type: none; }
            .toc a { text-decoration: none; color: #2980b9; }
            .footer { margin-top: 50px; text-align: center; color: #95a5a6; }
            
            /* Styling for internal links */
            a[href^="#"] { 
                color: #e74c3c; 
                text-decoration: none; 
                font-weight: 500; 
                border-bottom: 1px dotted #e74c3c;
                transition: all 0.2s ease;
            }
            a[href^="#"]:hover { 
                background-color: #fdf2f2; 
                text-decoration: none; 
                border-bottom: 1px solid #e74c3c;
                padding: 2px 4px;
                border-radius: 3px;
            }
        </style>
    </head>
    <body>
    """]
    
    # Header
    html_parts.append(f"""
        <h1>Glossário de Inteligência Artificial</h1>
        <div class="meta">
            <p><strong>Versão:</strong> {datetime.now().strftime('%Y-%m-%d')}</p>
            <p><strong>Total de categorias:</strong> {len(content)}</p>
            <p><strong>Gerado em:</strong> {datetime.now().strftime('%d/%m/%Y às %H:%M')}</p>
        </div>
    """)
    
    # Table of Contents
    html_parts.append('<div class="toc"><h2>Índice</h2><ul>')
    for category_name in content.keys():
        html_parts.append(f'<li><a href="#{category_name.lower().replace(" ", "-")}">{category_name}</a></li>')
    html_parts.append('</ul></div>')
    
    # Content
    md = markdown.Markdown(extensions=['extra', 'toc'])
    
    for category_name, terms in content.items():
        html_parts.append(f'<h2 id="{category_name.lower().replace(" ", "-")}">{category_name}</h2>')
        
        for term in terms:
            # Use file ID as anchor
            term_id = term['id']
            html_parts.append(f'<div class="term" id="{term_id}">')
            html_parts.append(f'<h3>{term["title"]}</h3>')
            
            # Process definition with internal links
            definition = term["definition"]
            definition = process_internal_links(definition, term_map)
            html_parts.append(f'<p><strong>{definition}</strong></p>')
            
            # Add explanation if exists, also with internal links
            if term["explanation"]:
                explanation = term["explanation"]
                explanation = process_internal_links(explanation, term_map)
                # Convert remaining markdown to HTML
                explanation = md.convert(explanation)
                html_parts.append(explanation)
            
            html_parts.append('</div>')
    
    # Footer
    html_parts.append(f"""
        <div class="footer">
            <hr>
            <p>Glossário de Inteligência Artificial - {datetime.now().year}</p>
            <p>Gerado automaticamente a partir do repositório: github.com/peninha/glossario-ia</p>
        </div>
    """)
    
    html_parts.append('</body></html>')
    
    return '\n'.join(html_parts)

def export_to_html(html_content, output_file='glossario_ia.html'):
    """Export HTML file"""
    try:
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(html_content)
        print(f"✅ HTML exportado: {output_file}")
        return True
    except Exception as e:
        print(f"❌ Erro ao exportar HTML: {e}")
        return False

def export_to_word(content, output_file='glossario_ia.docx'):
    """Export content to Word document"""
    try:
        from docx import Document
        from docx.shared import Inches
        
        doc = Document()
        total_terms = sum(len(terms) for terms in content.values())
        
        # Title
        title = doc.add_heading('Glossário de Inteligência Artificial', 0)
        
        # Metadata
        p = doc.add_paragraph()
        p.add_run(f'Total de termos: {total_terms}\n').bold = True
        p.add_run(f'Data: {datetime.now().strftime("%d/%m/%Y")}\n')
        
        doc.add_page_break()
        
        # Content
        for category_name, terms in content.items():
            doc.add_heading(category_name, level=1)
            
            for term in terms:
                # Term title
                doc.add_heading(term['title'], level=2)
                
                # Clean definition
                definition = term['definition']
                definition = re.sub(r'\*\*(.*?)\*\*', r'\1', definition)
                definition = re.sub(r'\*(.*?)\*', r'\1', definition)
                # For Word, we can keep links but make them simpler
                definition = re.sub(r'\[(.*?)\]\([^)]+\)', r'\1', definition)
                
                doc.add_paragraph(definition)
                
                # Add explanation if exists
                if term['explanation']:
                    explanation = term['explanation']
                    explanation = re.sub(r'\*\*(.*?)\*\*', r'\1', explanation)
                    explanation = re.sub(r'\*(.*?)\*', r'\1', explanation)
                    explanation = re.sub(r'\[(.*?)\]\([^)]+\)', r'\1', explanation)
                    
                    # Split explanation into paragraphs
                    for para in explanation.split('\n\n'):
                        if para.strip():
                            doc.add_paragraph(para.strip())
        
        doc.save(output_file)
        print(f"✅ Word exportado: {output_file}")
        return True
    except Exception as e:
        print(f"❌ Erro ao exportar Word: {e}")
        return False

def export_to_txt(content, output_file='glossario_ia.txt'):
    """Export content to plain text"""
    try:
        total_terms = sum(len(terms) for terms in content.values())
        
        lines = [
            "GLOSSÁRIO DE INTELIGÊNCIA ARTIFICIAL",
            "=" * 50,
            f"Total de termos: {total_terms}",
            f"Data: {datetime.now().strftime('%d/%m/%Y')}",
            "",
            "=" * 50,
            ""
        ]
        
        for category_name, terms in content.items():
            lines.append(f"{category_name.upper()}")
            lines.append("=" * len(category_name))
            lines.append("")
            
            for term in terms:
                lines.append(f"{term['title']}")
                lines.append("-" * len(term['title']))
                lines.append("")
                
                # Clean definition further
                definition = term['definition']
                # Remove markdown formatting
                definition = re.sub(r'\*\*(.*?)\*\*', r'\1', definition)  # Bold
                definition = re.sub(r'\*(.*?)\*', r'\1', definition)      # Italic
                definition = re.sub(r'\[(.*?)\]\([^)]+\)', r'\1', definition)  # Links
                
                lines.append(definition)
                
                if term['explanation']:
                    lines.append("")
                    explanation = term['explanation']
                    # Clean explanation
                    explanation = re.sub(r'\*\*(.*?)\*\*', r'\1', explanation)
                    explanation = re.sub(r'\*(.*?)\*', r'\1', explanation)
                    explanation = re.sub(r'\[(.*?)\]\([^)]+\)', r'\1', explanation)
                    lines.append(explanation)
                
                lines.append("")
                lines.append("")
        
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write('\n'.join(lines))
        
        print(f"✅ TXT exportado: {output_file}")
        return True
    except Exception as e:
        print(f"❌ Erro ao exportar TXT: {e}")
        return False

def main():
    """Main export function"""
    print("📄 Exportando Glossário de IA...")
    print("=" * 50)
    
    # Load content from docs
    content = load_docs_content()
    
    if not content:
        print("❌ Nenhum conteúdo encontrado na pasta docs/")
        return
    
    total_terms = sum(len(terms) for terms in content.values())
    print(f"📊 {len(content)} categorias, {total_terms} termos encontrados")
    print()
    
    # Generate HTML
    html_content = generate_html(content)
    
    # Export options
    exports = {
        'HTML': lambda: export_to_html(html_content),
        'Word': lambda: export_to_word(content),
        'TXT': lambda: export_to_txt(content)
    }
    
    # Ask user which formats to export
    print("Formatos disponíveis:")
    for i, format_name in enumerate(exports.keys(), 1):
        print(f"{i}. {format_name}")
    print("4. Todos")
    
    choice = input("\nEscolha o formato (1-4): ").strip()
    
    if choice == "4":
        # Export all
        for name, func in exports.items():
            func()
    elif choice in ["1", "2", "3"]:
        format_names = list(exports.keys())
        chosen_format = format_names[int(choice) - 1]
        exports[chosen_format]()
    else:
        print("❌ Opção inválida")
        return
    
    print("\n✅ Exportação concluída!")

if __name__ == "__main__":
    main()
